import streamlit as st
import openai
import base64
import re
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE LA PÁGINA WEB ---
st.set_page_config(
    page_title="Plataforma de Redacción Académica", 
    page_icon="📝", 
    layout="centered"
)

# --- INYECCIÓN DE ESTILOS CSS AVANZADOS (Diseño UI/UX Premium) ---
st.markdown(f"""
    <style>
    /* Fondo global de la aplicación: Tono Arena/Hueso #E5D8C8 */
    .stApp {{
        background-color: #E5D8C8 !important;
    }}
    
    /* Contenedor principal del bloque de contenido */
    .main .block-container {{
        padding-top: 1.5rem;
        padding-bottom: 3rem;
        max-width: 740px;
    }}
    
    /* Textos globales en Marrón Chocolate Profundo */
    .stMarkdown p, p, span, label, li, ul {{
        color: #2B1B17 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}
    
    /* Forzar texto blanco puro dentro de las notificaciones flotantes (st.toast) */
    div[data-testid="stToast"] p, 
    div[data-testid="stToast"] span, 
    div[data-testid="stToast"] div {{
        color: #FFFFFF !important;
    }}
    
    /* Título principal con estilo Banner Dinámico */
    .header-banner {{
        background-color: #FAF8F5;
        border: 1px solid rgba(43, 27, 23, 0.15);
        border-radius: 16px;
        padding: 2.2rem;
        text-align: center;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.02);
        transition: all 0.3s ease;
    }}
    
    /* Variación sutil para el modo informe (Borde izquierdo destacado) */
    .banner-informe {{
        border-left: 6px solid #4A322B !important;
        background-color: #FDFBF7;
    }}
    
    h1 {{
        color: #2B1B17 !important;
        font-weight: 800 !important;
        margin-bottom: 0.5rem !important;
    }}
    
    /* Estilos personalizados para las pestañas nativas de Streamlit */
    div[data-testid="stTabs"] button {{
        background-color: #FAF8F5 !important;
        color: #2B1B17 !important;
        border: 1px solid rgba(43, 27, 23, 0.12) !important;
        border-radius: 10px 10px 0px 0px !important;
        padding: 0.6rem 1.5rem !important;
        font-weight: 600 !important;
        font-size: 1.05rem !important;
        margin-right: 4px !important;
    }}
    
    div[data-testid="stTabs"] button[aria-selected="true"] {{
        background-color: #2B1B17 !important;
        color: #FFFFFF !important;
        border: 1px solid #2B1B17 !important;
    }}
    
    div[data-testid="stTabs"] button[aria-selected="true"] p {{
        color: #FFFFFF !important;
    }}

    /* Etiquetas de los componentes nativos */
    div[data-testid="stWidgetLabel"] p {{
        color: #2B1B17 !important;
        font-weight: 600 !important;
        font-size: 1.05rem !important;
    }}
    
    /* Tarjetas y Contenedores Generales */
    .custom-card, div[data-testid="stForm"], .stTextArea, div[data-testid="stFileUploader"], div[data-testid="stTextInput"] {{
        background-color: #FAF8F5 !important;
        border: 1px solid rgba(43, 27, 23, 0.12) !important;
        padding: 1.8rem !important;
        border-radius: 16px !important;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.03) !important;
        margin-bottom: 1.5rem !important;
    }}
    
    /* Selector de radio de entrada optimizado */
    div[data-testid="stRadio"] {{
        background-color: #FAF8F5 !important;
        border: 1px solid rgba(43, 27, 23, 0.12) !important;
        padding: 0.8rem 1.4rem !important;
        border-radius: 16px !important;
        box-shadow: 0 4px 14px rgba(0, 0, 0, 0.02) !important;
        margin-bottom: 1.5rem !important;
    }}
    
    div[data-testid="stRadio"] > div {{
        gap: 0.5rem !important;
    }}
    
    /* Píldoras para los Temas Listos */
    .tema-badge {{
        background-color: #EAE2D5;
        border-left: 4px solid #2B1B17;
        padding: 0.8rem 1.2rem;
        border-radius: 6px;
        margin-bottom: 0.6rem;
        font-weight: 500;
        color: #2B1B17;
        font-size: 0.95rem;
    }}
    
    /* Borde estético a la imagen subida */
    div[data-testid="stImage"] img {{
        border: 3px solid #2B1B17 !important;
        border-radius: 12px !important;
        box-shadow: 0 6px 18px rgba(0, 0, 0, 0.08) !important;
    }}
    
    /* CONFIGURACIÓN DE BOTONES PREMIUM */
    div.stButton > button, div[data-testid="stFileUploader"] button {{
        background-color: #2B1B17 !important;
        color: #FFFFFF !important;
        border: 2px solid #2B1B17 !important;
        border-radius: 10px !important;
        padding: 0.65rem 1.5rem !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
        width: 100% !important;
        box-shadow: 0 4px 10px rgba(43, 27, 23, 0.15) !important;
    }}
    
    div.stButton > button p, div[data-testid="stFileUploader"] button p {{
        color: #FFFFFF !important;
    }}
    
    div.stButton > button:hover, div[data-testid="stFileUploader"] button:hover {{
        background-color: #4A322B !important;
        border-color: #4A322B !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 16px rgba(43, 27, 23, 0.25) !important;
    }}
    
    /* Botón de Descarga Destacado Pro (#A4D8FF) */
    div.stDownloadButton > button {{
        background-color: #A4D8FF !important;
        color: #101D33 !important;
        border: 2px solid #A4D8FF !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
        font-size: 1.05rem !important;
        width: 100% !important;
        padding: 0.8rem !important;
        box-shadow: 0 6px 20px rgba(164, 216, 255, 0.35) !important;
        transition: all 0.25s ease !important;
    }}
    div.stDownloadButton > button p {{
        color: #101D33 !important;
    }}
    div.stDownloadButton > button:hover {{
        background-color: #2B1B17 !important;
        border-color: #2B1B17 !important;
        transform: translateY(-2px) !important;
    }}
    div.stDownloadButton > button:hover p {{
        color: #FFFFFF !important;
    }}

    /* Área interactiva del cargador de archivos */
    div[data-testid="stFileUploader"] section {{
        background-color: #FFFFFF !important;
        border: 2px dashed rgba(43, 27, 23, 0.25) !important;
        border-radius: 10px;
        padding: 1.2rem !important;
    }}
    </style>
""", unsafe_allow_html=True)

# --- CONTROL DE ACCESO SIMPLE & DESENCRIPCIÓN SECRETA ---
def check_access():
    """Verifica la contraseña y carga la API Key de forma segura desde Base64."""
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if "user_api_key" not in st.session_state:
        st.session_state["user_api_key"] = None

    if st.session_state["authenticated"] and st.session_state["user_api_key"]:
        return True

    st.markdown('''
        <div class="header-banner">
            <h1>🔒 Acceso Privado</h1>
            <p style="font-size: 1rem; margin: 0; color: #2B1B17; opacity: 0.85;">
                Introduce la contraseña de la aplicación para desbloquear la herramienta.
            </p>
        </div>
    ''', unsafe_allow_html=True)
    
    app_password = st.text_input("Contraseña de la Aplicación:", type="password", placeholder="Escribe la clave de la app...")
    
    if st.button("🔑 Entrar"):
        if app_password == "12345678":
            try:
                # -------------------------------------------------------------------------
                # REEMPLAZA EL TEXTO DE ABAJO POR TU CLAVE DE OPENAI EN BASE64
                # -------------------------------------------------------------------------
                key_ofuscada = "c2stcHJvai1BZlFuUnlnd2NQLTYwNEdvSEl6ZUFnRU12RXF4azNKQjVMQkdzeUNRbnc1OC0tNG5xZEhaUGVfdmZFSXh5MlR2T0pPX1ZsVzN6UlQzQmxia0ZKdWs5eWdsX2JTUE14ajBCaHhZcndxZi05T0k4UWdQQTN4d05RbFg2THdUVXNDdlA3LTNYc3RNcUt5Yll5Qlo5S01OMk9RNzRxd0E=" 
                
                key_descifrada = base64.b64decode(key_ofuscada).decode('utf-8')
                st.session_state["authenticated"] = True
                st.session_state["user_api_key"] = key_descifrada
                st.rerun()
            except Exception:
                st.error("❌ Ocurrió un error interno al validar el acceso.")
        else:
            st.error("❌ Contraseña de aplicación incorrecta.")
            
    return False

# La aplicación corre únicamente si pasa el control interactivo
if check_access():

    client = openai.OpenAI(api_key=st.session_state["user_api_key"])

    # --- NAVEGACIÓN SUPERIOR POR PESTAÑAS (BOTONES DE MODO) ---
    tab_investigacion, tab_informe = st.tabs(["🔬 Modo Investigación", "📊 Modo Informe Complejo"])

    def corregir_capitales_y_ortografia(texto):
        reemplazos = {
            r"\bgomez\b": "Gómez", r"\bperez\b": "Pérez", r"\bjimenez\b": "Jiménez",
            r"\bjuan\b": "Juan", r"\bvicente\b": "Vicente", r"\bmarcos\b": "Marcos",
        }
        for patron, reemplazo in reemplazos.items():
            texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
        texto = re.sub(r'(^[a-z]|(?<=\.\s)[a-z])', lambda m: m.group(1).upper(), texto)
        return texto

    # Lógica compartida para procesar y renderizar contenido según la pestaña activa
    def ejecutar_generador(modo_activo):
        temas_extraidos = []
        
        st.markdown('<div style="font-size: 1.05rem; font-weight: 700; color: #2B1B17; margin-bottom: 0.3rem;">🛠️ Método de entrada</div>', unsafe_allow_html=True)
        opcion = st.radio(f"Entrada ({modo_activo}):", ("Mediante una Imagen Scan", "Mediante Texto Manual"), label_visibility="collapsed", key=f"radio_{modo_activo}")

        if opcion == "Mediante una Imagen Scan":
            uploaded_file = st.file_uploader("Arrastra o selecciona la captura de tu asignación:", type=["jpg", "jpeg", "png"], key=f"file_{modo_activo}")
            if uploaded_file is not None:
                st.image(uploaded_file, caption="📸 Vista previa de la captura", use_container_width=True)
                
                if st.button("✨ Analizar y Estructurar Asignación", key=f"btn_scan_{modo_activo}"):
                    with st.spinner("La IA está leyendo los conceptos y organizando los temas..."):
                        try:
                            base64_image = base64.b64encode(uploaded_file.read()).decode('utf-8')
                            prompt_vision = (
                                "Analiza detalladamente esta imagen de una asignación escolar o universitaria.\n\n"
                                "REGLAS CRÍTICAS DE EXTRACCIÓN:\n"
                                "1. Descarta por completo cualquier dato administrativo, fechas de entrega, ponderaciones, "
                                "modalidades, palabras como 'defensas', 'informe escrito', 'evaluación', o nombres de materias.\n"
                                "2. Identifica cuál es el TEMA central de la hoja (por ejemplo, en este caso es 'El Turismo en Venezuela').\n"
                                "3. Toma cada uno de los subpuntos con asterisco (*concepto, *caracteristicas, *impacto economico, *problematica, *situacion actual, *turismo regional) "
                                "y fusiónalos lógicamente con el tema central.\n\n"
                                "Devuelve exactamente esta lista separada por comas, sin preámbulos: El Turismo en Venezuela: concepto, Características del turismo en Venezuela, Impacto económico del turismo en Venezuela, Problemáticas del turismo en Venezuela, Situación actual del turismo en Venezuela, Turismo regional en Venezuela"
                            )
                            
                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": [
                                    {"type": "text", "text": prompt_vision},
                                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                                ]}]
                            )
                            contenido = response.choices[0].message.content
                            for etiqueta in ["Tema principal:", "Tema global:", "Puntos:", "Subpuntos:", "Temas:"]:
                                contenido = contenido.replace(etiqueta, "")
                            
                            st.session_state[f'temas_{modo_activo}'] = [tema.strip() for tema in contenido.split(',') if tema.strip()]
                            st.toast("¡Temas leídos de forma exitosa!", icon="🚀")
                        except Exception as e:
                            st.error(f"Error analizando la imagen: {e}")

        elif opcion == "Mediante Texto Manual":
            entrada = st.text_area("Pega o escribe la lista de temas (separados por comas):", 
                                   placeholder="Ej: Concepto y características de las células, Leyes de Mendel en la genética moderna", height=120, key=f"txt_{modo_activo}")
            if st.button("⚙️ Cargar Temas", key=f"btn_text_{modo_activo}"):
                if entrada:
                    st.session_state[f'temas_{modo_activo}'] = [tema.strip() for tema in entrada.split(',') if tema.strip()]
                    st.toast("Temas guardados.", icon="✅")
                else:
                    st.warning("El campo de texto está vacío.")

        # Procesamiento final de los documentos
        key_estado = f'temas_{modo_activo}'
        if key_estado in st.session_state:
            temas_extraidos = st.session_state[key_estado]
            palabras_prohibidas = ["4to año", "lunes", "martes", "miércoles", "jueves", "viernes", "investigación", "entrega", "individual", "ponderación"]
            
            temas_finales = []
            for t in temas_extraidos:
                t_limpio = re.sub(r'^\d+[\s\.\:\-\"\']*', '', t).strip(' "\'')
                if t_limpio and not any(p in t_limpio.lower() for p in palabras_prohibidas):
                    temas_finales.append(t_limpio)

            if temas_finales:
                st.markdown('<div class="custom-card">', unsafe_allow_html=True)
                st.markdown("<h3 style='margin-bottom: 1rem;'>📋 Puntos identificados para desarrollo:</h3>", unsafe_allow_html=True)
                for item in temas_finales:
                    st.markdown(f'<div class="tema-badge">{item}</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

                if st.button(f"🚀 Generar Documento ({modo_activo})", key=f"btn_run_{modo_activo}"):
                    doc = Document()
                    progreso = st.progress(0)
                    status_text = st.empty()

                    for i, tema in enumerate(temas_finales):
                        numero = i + 1
                        status_text.markdown(f"⏳ **Procesando [{numero}/{len(temas_finales)}]:** *{tema}*")
                        
                        t_clean = tema.replace("-", "-").replace("•", "").strip()
                        t_clean = corregir_capitales_y_ortografia(t_clean.strip('¿?'))
                        if len(t_clean) > 0:
                            t_clean = t_clean[0].upper() + t_clean[1:]
                        
                        t_low = t_clean.lower()
                        es_pregunta = t_low.startswith(("que", "como", "cual", "por que", "por qué", "quien", "quién", "donde", "cómo", "qué", "cuál"))
                        titulo_final = f"{numero}. ¿{t_clean}?" if es_pregunta else f"{numero}. {t_clean}"

                        # ASIGNACIÓN DE PROMPTS SEGÚN LA PESTAÑA SELECCIONADA (CON FILTROS ANTI-DESVÍO)
                        if modo_activo == "Investigacion":
                            es_tema_importante = any(palabra in tema.lower() for palabra in ["gobierno", "transicion", "transición", "impacto", "situacion", "situación", "leyes", "teoria"])
                            rango_palabras = "entre 160 y 180 palabras" if es_tema_importante else "entre 115 y 125 palabras"
                            formato_lista = " DEBES incluir una lista formal estructurada con viñetas 'Componente: descripción'." if i % 2 == 0 else " No uses listas, redacta completamente en párrafos continuos y fluidos."
                            
                            instrucciones_redaccion = (
                                f"Eres un académico e investigador experto. Desarrolla el punto solicitado con rigurosidad.\n"
                                f"REGLA CRÍTICA DE CONTEXTO: Tu respuesta debe centrarse TOTAL Y EXCLUSIVAMENTE en Venezuela y su relación directa con el título '{tema}'. Está prohibido hablar de otros países o continentes de forma genérica.\n"
                                f"REGLA CRÍTICA DE EXTENSIÓN: El texto completo debe tener obligatoriamente {rango_palabras}. Mantén un equilibrio preciso.\n"
                                f"No incluyas títulos en tu respuesta. Prohibido usar el signo de punto y coma (;). No repitas ideas ni frases.\n"
                                f"MINÚSCULAS Y MAYÚSCULAS: Todo el texto regular debe ir en minúsculas, EXCEPTO la primera letra de cada oración y nombres propios.\n"
                                f"LISTAS: Usa el símbolo •. El formato obligatorio de cada punto debe ser 'Componente: descripción breve' (máximo 20 palabras por punto).{formato_lista}"
                            )
                        else:
                            # MODULO INFORME MEJORADO: Estable, equilibrado y amarrado estrictamente al contexto geográfico local
                            instrucciones_redaccion = (
                                f"Eres un analista de datos e historiador experto. Desarrolla un informe corporativo de alta calidad sobre el punto específico: '{tema}'.\n"
                                f"REGLA CRÍTICA DE CONTEXTO GEOGRÁFICO: Todo el contenido debe desarrollarse dentro del marco geográfico e histórico de VENEZUELA. Queda prohibido desvilarse a hablar de la Amazonía de Brasil, crisis globales de EE.UU. o el cambio climático general de forma aislada. Si el tema es impacto o problemática, habla de la realidad venezolana real (ej. la crisis del sector hotelero local, la devaluación, la falta de vuelos, servicios básicos, etc.).\n"
                                f"REGLA DE EXTENSIÓN ESTABLE: Todo el texto de este punto debe tener un desarrollo equilibrado de entre 200 y 240 palabras (ni más, ni menos). Asegura un largo uniforme en todos los puntos.\n"
                                f"REGLAS ESTRUCTURALES:\n"
                                f"1. Incluye datos numéricos, porcentajes o estimaciones realistas aplicados exclusivamente a Venezuela.\n"
                                f"2. Menciona acontecimientos clave locales con sus años o períodos históricos respectivos (ej. Ley de Turismo, hitos de los años 90, crisis post-2014, etc.).\n"
                                f"3. Prohibido terminantemente repetir párrafos, oraciones o ideas. Cada línea debe aportar valor nuevo.\n"
                                f"4. Prohibido usar el signo de punto y coma (;).\n"
                                f"5. Organiza la información combinando un párrafo narrativo fluido con un bloque analítico de viñetas (•) que detalle hechos específicos."
                            )

                        try:
                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {"role": "system", "content": instrucciones_redaccion},
                                    {"role": "user", "content": f"Desarrolle el contenido detallado y equilibrado exclusivamente para el punto: '{tema}' sin repetir información."}
                                ],
                                temperature=0.2, # Bajamos la temperatura para evitar redundancias y bucles
                                frequency_penalty=1.0 # Penalización para asegurar que no repita palabras ni frases idénticas
                            )
                            
                            texto_generado = response.choices[0].message.content.replace("*", "").replace(";", ".")
                            texto_generado = corregir_capitales_y_ortografia(texto_generado)

                            # --- CONSTRUCCIÓN DEL WORD ---
                            h = doc.add_paragraph()
                            run_h = h.add_run(titulo_final)
                            run_h.bold = True
                            run_h.font.name = 'Canva Sans'
                            run_h.font.size = Pt(15)

                            lineas = texto_generado.split('\n')
                            for linea in lineas:
                                linea = linea.strip()
                                if not linea: continue
                                
                                if not linea.startswith("•") and len(linea.split()) > 65:
                                    puntos = linea.split('. ')
                                    mitad = len(puntos) // 2
                                    bloques = [". ".join(puntos[:mitad]) + ".", ". ".join(puntos[mitad:])] if mitad > 0 else [linea]
                                else:
                                    bloques = [linea]

                                for bloque in bloques:
                                    p = doc.add_paragraph()
                                    if bloque.strip().startswith("•"):
                                        p.add_run("• ").bold = True
                                        cont = bloque.strip().lstrip("• ").strip()
                                        if ":" in cont:
                                            sub, desc = cont.split(":", 1)
                                            run_sub = p.add_run(f"{sub.strip()}:")
                                            run_sub.bold = True
                                            p.add_run(desc)
                                        else:
                                            p.add_run(cont)
                                        p.paragraph_format.left_indent = Pt(24)
                                    else:
                                        p.add_run(bloque.strip())
                                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    
                                    for run in p.runs:
                                        run.font.name = 'Canva Sans'
                                        run.font.size = Pt(13)

                            doc.add_paragraph()
                        except Exception as e:
                            st.error(f"Error generando el tema {tema}: {e}")
                        
                        progreso.progress(int((numero / len(temas_finales)) * 100))

                    status_text.success("🎉 ¡El informe ha sido completado con éxito total!")
                    
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)

                    st.markdown('<div style="margin-top: 1rem;"></div>', unsafe_allow_html=True)
                    st.download_button(
                        label="💾 DESCARGAR INFORME EN WORD (.DOCX)",
                        data=bio,
                        file_name=f"Asignacion_{modo_activo}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{modo_activo}"
                    )
            else:
                st.warning("No se detectaron temas válidos para redactar.")

    # --- PESTAÑA 1: MODO INVESTIGACIÓN ---
    with tab_investigacion:
        st.markdown('''
            <div class="header-banner">
                <h1>🔬 Módulo de Investigación Académica</h1>
                <p style="font-size: 1.1rem; font-weight: 500; margin: 0; color: #2B1B17; opacity: 0.85;">
                    Genera respuestas analíticas, compactas y fluidas. Ideal para cuestionarios, tareas dirigidas y explicaciones conceptuales directas.
                </p>
            </div>
        ''', unsafe_allow_html=True)
        ejecutar_generador("Investigacion")

    # --- PESTAÑA 2: MODO INFORME COMPLEJO ---
    with tab_informe:
        st.markdown('''
            <div class="header-banner banner-informe">
                <h1 style="color: #4A322B !important;">📊 Módulo de Informes Complejos</h1>
                <p style="font-size: 1.1rem; font-weight: 500; margin: 0; color: #2B1B17; opacity: 0.85;">
                    Redacción corporativa de largo alcance vinculada estrictamente al contexto nacional. Incluye de forma equilibrada <b>métricas, fechas críticas y acontecimientos históricos locales</b>.
                </p>
            </div>
        ''', unsafe_allow_html=True)
        ejecutar_generador("Informe")